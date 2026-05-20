"""Chargement de documents: PDF, DOCX, TXT, MD, HTML, URL."""
import logging
import os
from pathlib import Path
from typing import List, Dict, Any
from dataclasses import dataclass
import urllib.request
import urllib.parse

# Tags HTML à supprimer avant extraction du texte.
# Défini une seule fois, partagé par _load_html et scrape_url.
_HTML_NOISE_TAGS = [
    "script", "style", "nav", "footer", "header",
    "aside", "noscript", "iframe", "form", "button",
    "select", "input", "textarea", "figure", "figcaption",
    "meta", "link",
]

logger = logging.getLogger(__name__)


@dataclass
class Document:
    content: str
    metadata: Dict[str, Any]

def _load_excel(path: str) -> str:
    """Charge un fichier Excel — fallback texte pour les formats simples."""
    docs = load_excel_as_documents(path)
    return "\n\n".join(d.content for d in docs)


def _detect_header_row(ws, max_scan: int = 20) -> int:
    """Détecte la ligne d'en-tête parmi les max_scan premières lignes.

    Stratégie : retourne la ligne avec le PLUS de cellules texte non-numériques.
    Si plusieurs lignes sont à égalité, retourne la première (titre avant header).
    Un titre occupe souvent 1 cellule fusionnée ; le header remplit toutes les colonnes.
    """
    best_row   = 1
    best_count = 0
    for row_idx, row in enumerate(ws.iter_rows(min_row=1, max_row=max_scan, values_only=True), 1):
        text_cells = 0
        for v in row:
            if v is None:
                continue
            s = str(v).strip()
            if len(s) < 2:
                continue
            try:
                float(s.replace(",", "."))
            except ValueError:
                text_cells += 1
        # Strictement supérieur : préfère la première ligne maximale
        if text_cells > best_count:
            best_count = text_cells
            best_row   = row_idx
    return best_row if best_count >= 2 else 1


def load_excel_as_documents(path: str) -> List[Document]:
    """Charge un fichier Excel et retourne UN Document par ligne.
    Chaque entrée est un texte autonome avec les noms de colonnes,
    ce qui permet au retriever de les associer sémantiquement
    et au chunker de ne pas les couper arbitrairement."""
    import openpyxl
    p = Path(path)
    wb = openpyxl.load_workbook(path, data_only=True)
    docs = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        # Auto-détection de la ligne d'en-tête (gère les Excel avec titre/métadonnées avant)
        header_row = _detect_header_row(ws)
        headers = []
        for cell in ws[header_row]:
            headers.append(str(cell.value).strip() if cell.value else "")

        for row_idx, row in enumerate(ws.iter_rows(min_row=header_row + 1, values_only=True), header_row + 1):
            values = [str(v).strip() if v is not None else "" for v in row]
            if not any(values):
                continue
            # Construire un texte structuré "Colonne: Valeur" par ligne
            entry_parts = []
            for header, value in zip(headers, values):
                if header and value:
                    entry_parts.append(f"{header}: {value}")
            if not entry_parts:
                continue
            # Préfixer avec le nom du fichier source (sans extension)
            # pour aider le retriever et le LLM à distinguer les sources
            source_label = p.stem.upper()  # DIRECTION, DEPARTEMENT, SERVICE, POSTE
            content = f"[{source_label}] " + " | ".join(entry_parts)
            docs.append(Document(
                content=content,
                metadata={
                    "source": str(p),
                    "filename": p.name,
                    "extension": p.suffix.lower(),
                    "sheet": sheet_name,
                    "row": row_idx,
                    "size_bytes": p.stat().st_size,
                },
            ))
    return docs

def _load_pdf(path: str) -> str:
    from pypdf import PdfReader
    reader = PdfReader(path)
    pages = [page.extract_text() or "" for page in reader.pages]
    text  = "\n\n".join(p for p in pages if p.strip())

    if not text.strip():
        # PDF scanné (image uniquement) → tentative OCR
        try:
            import pytesseract
            from pdf2image import convert_from_path
            logger.info("PDF '%s' : texte vide, tentative OCR (pytesseract)...", path)
            images = convert_from_path(path, dpi=200)
            # Langue configurable via OCR_LANG (ex: "fra+eng+ara"). Défaut: "fra+eng".
            ocr_lang = os.environ.get("OCR_LANG", "fra+eng")
            ocr_pages = []
            for img in images:
                page_text = pytesseract.image_to_string(img, lang=ocr_lang)
                if page_text.strip():
                    ocr_pages.append(page_text.strip())
            if ocr_pages:
                text = "\n\n".join(ocr_pages)
                logger.info("  ✓ OCR réussi : %d page(s), %d chars", len(ocr_pages), len(text))
            else:
                logger.warning("  ✗ OCR : aucun texte extrait de '%s'", path)
        except ImportError:
            logger.warning(
                "PDF '%s' vide et OCR non disponible. "
                "Pour les PDF scannés : pip install pytesseract pdf2image "
                "(+ installer Tesseract-OCR sur le système).", path
            )
        except Exception as e:
            logger.warning("  ✗ OCR échoué pour '%s' : %s", path, e)

    return text


_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

def _load_docx(path: str) -> str:
    """Extrait le texte d'un DOCX en préservant l'ordre paragraphes + tableaux."""
    from docx import Document as DocxDoc
    doc = DocxDoc(path)
    parts = []
    for child in doc.element.body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            text = "".join(n.text or "" for n in child.iter() if n.text)
            if text.strip():
                parts.append(text.strip())
        elif tag == "tbl":
            # Convertit chaque ligne du tableau en "col1 | col2 | col3"
            for tr in child.iter(f"{{{_W_NS}}}tr"):
                cells = []
                for tc in tr.findall(f"{{{_W_NS}}}tc"):
                    cell_text = "".join(n.text or "" for n in tc.iter() if n.text)
                    if cell_text.strip():
                        cells.append(cell_text.strip())
                if cells:
                    parts.append(" | ".join(cells))
    return "\n".join(parts)


def _load_html(path: str) -> str:
    from bs4 import BeautifulSoup
    with open(path, encoding="utf-8") as f:
        soup = BeautifulSoup(f.read(), "html.parser")
    for tag in soup(_HTML_NOISE_TAGS):
        tag.decompose()
    return soup.get_text(separator="\n", strip=True)


def _load_text(path: str) -> str:
    with open(path, encoding="utf-8", errors="replace") as f:
        return f.read()


LOADERS = {
    ".pdf": _load_pdf,
    ".docx": _load_docx,
    ".doc": _load_docx,
    ".html": _load_html,
    ".htm": _load_html,
    ".txt": _load_text,
    ".md": _load_text,
    ".markdown": _load_text,
    ".rst": _load_text,
    ".csv": _load_text,
    ".json": _load_text,
      ".xlsx":     _load_excel,   # ← ajouté
    ".xls":      _load_excel,   # ← ajouté
}


def load_document(path: str) -> Document:
    p = Path(path)
    ext = p.suffix.lower()
    if ext not in LOADERS:
        raise ValueError(f"Format non supporté: {ext}")

    content = LOADERS[ext](str(p))
    return Document(
        content=content,
        metadata={
            "source": str(p),
            "filename": p.name,
            "extension": ext,
            "size_bytes": p.stat().st_size,
        },
    )


def scrape_url(url: str, timeout: int = 15) -> Document:
    """Scrape une URL HTTP/HTTPS et retourne un Document."""
    parsed = urllib.parse.urlparse(url)
    if parsed.scheme not in ("http", "https"):
        raise ValueError(f"URL invalide (http/https requis): {url}")

    req = urllib.request.Request(url, headers={"User-Agent": "RAGBot/1.0"})
    with urllib.request.urlopen(req, timeout=timeout) as resp:
        raw_html = resp.read()

    from bs4 import BeautifulSoup
    soup = BeautifulSoup(raw_html, "html.parser")
    for tag in soup(_HTML_NOISE_TAGS):
        tag.decompose()

    title = soup.title.string.strip() if soup.title and soup.title.string else parsed.netloc
    content = soup.get_text(separator="\n", strip=True)

    # Nom de fichier safe pour les IDs ChromaDB
    safe_name = urllib.parse.quote(url, safe="").replace("%", "_")[:80] + ".url"

    return Document(
        content=content,
        metadata={
            "source": url,
            "filename": safe_name,
            "extension": ".url",
            "title": title,
        },
    )


def load_directory(directory: str) -> List[Document]:
    docs = []
    excel_exts = {".xlsx", ".xls"}
    for file in sorted(Path(directory).rglob("*")):
        if not file.is_file():
            continue
        ext = file.suffix.lower()
        if ext not in LOADERS:
            continue
        try:
            if ext in excel_exts:
                # Excel : un Document par ligne pour un meilleur retrieval
                excel_docs = load_excel_as_documents(str(file))
                docs.extend(excel_docs)
                logger.info("  Chargé: %s (%d entrées Excel)", file.name, len(excel_docs))
            else:
                doc = load_document(str(file))
                if doc.content.strip():
                    docs.append(doc)
                    logger.info("  Chargé: %s (%d chars)", file.name, len(doc.content))
        except Exception as e:
            logger.warning("  Ignoré %s: %s", file.name, e)
    return docs
